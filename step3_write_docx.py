"""
step3_write_docx.py
Reads data/jobs_scored.json (written by Claude Code in Step 3 of CLAUDE.md)
and generates a tailored .docx resume for each qualified job.

Template placeholders in your .docx files:
  [[SUMMARY]]           — replaced with tailored_summary
  [[KPI_BULLET_1..7]]   — replaced with kpi_bullets[0..6]
  [[FACTSET_BULLET_1..6]] — replaced with factset_bullets[0..5]
"""

import json
import re
from pathlib import Path
from datetime import datetime

from docx import Document

IN      = Path("data/jobs_scored.json")
OUT_DIR = Path("output/resumes")
OUT_DIR.mkdir(parents=True, exist_ok=True)


def run():
    if not IN.exists():
        print(f"⚠ {IN} not found — skipping docx generation")
        return 0

    data = json.loads(IN.read_text())
    jobs = data.get("qualified_jobs", [])

    if not jobs:
        print("No qualified jobs — nothing to write.")
        return 0

    written = 0
    for job in jobs:
        path = _write_one(job)
        if path:
            job["resume_file"] = str(path)
            written += 1

    # Write back with resume_file paths filled in
    IN.write_text(json.dumps(data, indent=2, default=str))
    print(f"✓ Wrote {written} resume files to {OUT_DIR}/")
    return written


def _write_one(job: dict) -> Path | None:
    emp_type = job.get("employment_type", "full_time")
    template = _get_template(emp_type)

    try:
        doc = Document(str(template))
    except Exception as e:
        print(f"  ⚠ Could not open template {template}: {e}")
        return None

    replacements = _build_replacements(job)

    for para in doc.paragraphs:
        _replace(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace(para, replacements)

    company  = re.sub(r"[^\w\s-]", "", job.get("company", "Co")).strip().replace(" ", "_")
    title    = re.sub(r"[^\w\s-]", "", job.get("title", "Role")).strip().replace(" ", "_")[:28]
    et       = emp_type[:3].upper()
    date_str = datetime.now().strftime("%Y%m%d")
    fname    = f"Gautham_{company}_{title}_{et}_{date_str}.docx"
    out_path = OUT_DIR / fname

    doc.save(str(out_path))
    print(f"  ✓ {fname}")
    return out_path


def _build_replacements(job: dict) -> dict:
    r = {"[[SUMMARY]]": job.get("tailored_summary", "")}
    for i, b in enumerate(job.get("kpi_bullets", []), 1):
        r[f"[[KPI_BULLET_{i}]]"] = b
    for i, b in enumerate(job.get("factset_bullets", []), 1):
        r[f"[[FACTSET_BULLET_{i}]]"] = b
    return r


def _replace(para, replacements: dict):
    for old, new in replacements.items():
        if old in para.text:
            for run in para.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    return
            # Fallback: rebuild across runs
            full = para.text
            if old in full:
                new_text = full.replace(old, new)
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = new_text
                else:
                    para.add_run(new_text)


def _get_template(emp_type: str) -> Path:
    primary = Path(f"templates/resume_{'c2c' if emp_type == 'c2c' else 'ft'}.docx")
    fallback = Path("templates/resume_ft.docx")
    if primary.exists():
        return primary
    if fallback.exists():
        return fallback
    return _make_basic_template(emp_type)


def _make_basic_template(emp_type: str) -> Path:
    """Generates a minimal template when none exists — replace with your own."""
    p = Path(f"templates/resume_{emp_type}.docx")
    p.parent.mkdir(exist_ok=True)
    doc = Document()
    doc.add_heading("GAUTHAM POTHANA", 0)
    doc.add_paragraph("Java Full Stack Developer | Backend & Cloud Engineer")
    doc.add_paragraph("Cincinnati, OH | gauthampothana007@gmail.com | 513-918-9550")
    doc.add_heading("PROFESSIONAL SUMMARY", 1)
    doc.add_paragraph("[[SUMMARY]]")
    doc.add_heading("EXPERIENCE", 1)
    doc.add_paragraph("KPI Solutions — Cincinnati, OH | Dec 2023 – Present")
    doc.add_paragraph("Java Full Stack Developer")
    for i in range(1, 8):
        doc.add_paragraph(f"[[KPI_BULLET_{i}]]", style="List Bullet")
    doc.add_paragraph("FactSet — Hyderabad, India | Jan 2020 – Jul 2022")
    doc.add_paragraph("Software Engineer I")
    for i in range(1, 7):
        doc.add_paragraph(f"[[FACTSET_BULLET_{i}]]", style="List Bullet")
    doc.save(str(p))
    print(f"  Generated basic template: {p}")
    return p


if __name__ == "__main__":
    n = run()
    print(f"Done — {n} resumes written.")
